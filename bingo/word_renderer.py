# General Bingo Generator
# Copyright (c) 2026 Craebs Media
# Licensed under the PolyForm Noncommercial License 1.0.0
#
# Author: Craebs Media
#
# DOCX renderer. This module has no GUI dependencies: it receives finished cards
# and layout settings and creates a Word document with python-docx.

from __future__ import annotations

from pathlib import Path

import docx
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt

from .layout_engine import calculate_layout_metrics, estimate_header_height
from .models import CornerImage, GameSettings
from .word_utils import (
    rgb_from_hex,
    set_cell_appearance,
    set_cell_margins,
    set_column_width,
    set_row_height_exact,
    set_table_borders_none,
    set_table_layout_fixed,
)

TWIPS_PER_INCH = 1440
PROJECT_ROOT = Path(__file__).resolve().parents[1]


class WordBingoRenderer:
    """Render a list of Bingo cards into a .docx file."""

    def __init__(self, settings: GameSettings):
        self.settings = settings
        self.metrics = calculate_layout_metrics(settings)

    def render(self, cards: list[list[str]], output_path: str | Path | None = None) -> Path:
        """Create and save the Word document."""
        output = Path(output_path or self.settings.output_path)
        output.parent.mkdir(parents=True, exist_ok=True)

        doc = docx.Document()
        self._setup_sections(doc)
        self._setup_footer(doc)

        for card_index, card_items in enumerate(cards):
            self._add_card_page(doc, card_items)
            if card_index < len(cards) - 1:
                doc.add_page_break()

        doc.save(output)
        return output

    def _setup_sections(self, doc: docx.Document) -> None:
        """Apply page size, margins and footer distance to every section."""
        for section in doc.sections:
            section.page_width = Inches(self.settings.page_width_inch)
            section.page_height = Inches(self.settings.page_height_inch)
            section.top_margin = Inches(self.settings.margins.top)
            section.bottom_margin = Inches(self.settings.margins.bottom)
            section.left_margin = Inches(self.settings.margins.left)
            section.right_margin = Inches(self.settings.margins.right)
            section.footer_distance = Inches(self.settings.margins.footer_distance)

    def _usable_width_inch(self) -> float:
        return self.settings.page_width_inch - self.settings.margins.left - self.settings.margins.right

    def _resolve_image_path(self, image: CornerImage) -> Path | None:
        """Accept absolute paths and paths relative to the project root."""
        if not image.path:
            return None
        raw = Path(image.path)
        candidates = [raw]
        if not raw.is_absolute():
            candidates.append(PROJECT_ROOT / raw)
        for candidate in candidates:
            if candidate.exists():
                return candidate
        return None

    def _add_image_if_present(self, paragraph, image: CornerImage) -> None:
        """Add an image to a paragraph if the file exists."""
        path = self._resolve_image_path(image)
        if path:
            paragraph.add_run().add_picture(str(path), width=Inches(image.width_inch))

    @staticmethod
    def _clear_container(container) -> None:
        """Remove all existing content from a header/footer container.

        This avoids duplicated footer images when the document is regenerated or
        when python-docx creates a default empty footer paragraph.
        """
        element = container._element
        for child in list(element):
            element.remove(child)

    def _setup_footer(self, doc: docx.Document) -> None:
        """Place bottom-left and bottom-right images in the Word footer.

        Footer images do not consume space in the normal document flow. The grid
        position is calculated separately in layout_engine.py with enough footer
        reserve to avoid collisions.
        """
        for section in doc.sections:
            section.different_first_page_header_footer = False
            footer = section.footer
            self._clear_container(footer)

            table = footer.add_table(rows=1, cols=3, width=Inches(self._usable_width_inch()))
            table.alignment = WD_TABLE_ALIGNMENT.CENTER
            table.allow_autofit = False
            set_table_layout_fixed(table)
            set_table_borders_none(table)

            left_w = max(self.settings.bottom_left_image.width_inch + 0.1, 1.4)
            right_w = max(self.settings.bottom_right_image.width_inch + 0.1, 1.4)
            middle_w = max(self._usable_width_inch() - left_w - right_w, 0.5)
            widths = [left_w, middle_w, right_w]

            for i, width in enumerate(widths):
                table.columns[i].width = Inches(width)
                cell = table.rows[0].cells[i]
                cell.width = Inches(width)
                set_cell_margins(cell, 0, 0, 0, 0)
                cell.vertical_alignment = WD_ALIGN_VERTICAL.BOTTOM

            p_left = table.rows[0].cells[0].paragraphs[0]
            p_left.alignment = WD_ALIGN_PARAGRAPH.LEFT
            self._zero_paragraph_spacing(p_left)
            self._add_image_if_present(p_left, self.settings.bottom_left_image)

            p_right = table.rows[0].cells[2].paragraphs[0]
            p_right.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            self._zero_paragraph_spacing(p_right)
            self._add_image_if_present(p_right, self.settings.bottom_right_image)

    @staticmethod
    def _zero_paragraph_spacing(paragraph) -> None:
        paragraph.paragraph_format.space_before = Pt(0)
        paragraph.paragraph_format.space_after = Pt(0)
        paragraph.paragraph_format.line_spacing = 1

    def _add_card_page(self, doc: docx.Document, card_items: list[str]) -> None:
        self._add_header_area(doc)
        self._add_spacer_before_intro(doc)
        self._add_intro_text(doc)
        self._add_spacer_between_intro_and_grid(doc)
        self._add_bingo_grid(doc, card_items)

    def _add_header_area(self, doc: docx.Document) -> None:
        """Add top corner images and the centered title/subtitle/date block."""
        table = doc.add_table(rows=1, cols=3)
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        table.allow_autofit = False
        set_table_layout_fixed(table)
        set_table_borders_none(table)

        # Keep corner columns narrow so long titles get the full usable width.
        # The images may be slightly wider than the column; Word still keeps
        # them inside the page margins because the whole table is centered.
        corner_w = max(0.85, min(max(self.settings.top_left_image.width_inch, self.settings.top_right_image.width_inch), 1.05))
        middle_w = max(self._usable_width_inch() - 2 * corner_w, 2.0)
        widths = [corner_w, middle_w, corner_w]

        for i, width in enumerate(widths):
            table.columns[i].width = Inches(width)
            cell = table.rows[0].cells[i]
            cell.width = Inches(width)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.TOP
            set_cell_margins(cell, 0, 0, 0, 0)

        p_left = table.rows[0].cells[0].paragraphs[0]
        p_left.alignment = WD_ALIGN_PARAGRAPH.LEFT
        self._zero_paragraph_spacing(p_left)
        self._add_image_if_present(p_left, self.settings.top_left_image)

        p_right = table.rows[0].cells[2].paragraphs[0]
        p_right.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        self._zero_paragraph_spacing(p_right)
        self._add_image_if_present(p_right, self.settings.top_right_image)

        cell_text = table.rows[0].cells[1]
        p_title = cell_text.paragraphs[0]
        p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p_title.paragraph_format.space_before = Pt(self.settings.title_top_spacing_pt)
        p_title.paragraph_format.space_after = Pt(0)
        run_title = p_title.add_run(self.settings.title)
        run_title.font.name = self.settings.title_font
        run_title.font.size = Pt(self.settings.title_size_pt)
        run_title.font.bold = True
        run_title.font.color.rgb = rgb_from_hex(self.settings.title_color)

        p_sub = cell_text.add_paragraph()
        p_sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p_sub.paragraph_format.space_before = Pt(0)
        p_sub.paragraph_format.space_after = Pt(0)
        run_sub = p_sub.add_run(self.settings.subtitle)
        run_sub.font.name = self.settings.subtitle_font
        run_sub.font.size = Pt(self.settings.subtitle_size_pt)
        run_sub.font.italic = True
        run_sub.font.color.rgb = rgb_from_hex(self.settings.subtitle_color)

        p_date = cell_text.add_paragraph()
        p_date.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p_date.paragraph_format.space_before = Pt(0)
        p_date.paragraph_format.space_after = Pt(4)
        run_date = p_date.add_run(self.settings.date)
        run_date.font.name = self.settings.date_font
        run_date.font.size = Pt(self.settings.date_size_pt)
        run_date.font.italic = True
        run_date.font.color.rgb = rgb_from_hex(self.settings.date_color)

    def _add_spacer_before_intro(self, doc: docx.Document) -> None:
        """Move the intro line to the calculated position directly above the grid."""
        estimated_header = estimate_header_height(self.settings)
        gap_after_header = max(0.0, self.metrics.intro_top - self.settings.margins.top - estimated_header)
        spacer = doc.add_paragraph()
        spacer.paragraph_format.space_before = Pt(0)
        spacer.paragraph_format.space_after = Pt(gap_after_header * 72.0)

    def _add_intro_text(self, doc: docx.Document) -> None:
        """Add the intro text aligned to the grid width, not to the page title."""
        if not self.settings.intro_text.strip():
            return
        table = doc.add_table(rows=1, cols=1)
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        table.allow_autofit = False
        set_table_layout_fixed(table)
        set_table_borders_none(table)
        table.columns[0].width = Inches(self.metrics.grid_width)
        cell = table.rows[0].cells[0]
        cell.width = Inches(self.metrics.grid_width)
        set_cell_margins(cell, 0, 0, 0, 0)
        p = cell.paragraphs[0]
        p.alignment = self._paragraph_alignment(self.settings.intro_alignment)
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(0)
        run_intro = p.add_run(self.settings.intro_text)
        run_intro.font.name = self.settings.intro_font
        run_intro.font.size = Pt(self.settings.intro_size_pt)
        run_intro.font.bold = True
        run_intro.font.color.rgb = rgb_from_hex(self.settings.intro_color)

    def _add_spacer_between_intro_and_grid(self, doc: docx.Document) -> None:
        """Add an optional micro-gap between intro and grid.

        When the gap is 0, no paragraph is inserted. That prevents Word from
        creating a visible blank line between the intro table and the matrix.
        """
        gap = max(0.0, self.settings.grid_top_gap_inch)
        if gap <= 0.001:
            return
        spacer = doc.add_paragraph()
        spacer.paragraph_format.space_before = Pt(0)
        spacer.paragraph_format.space_after = Pt(gap * 72.0)

    @staticmethod
    def _paragraph_alignment(value: str):
        if value == "left":
            return WD_ALIGN_PARAGRAPH.LEFT
        if value == "right":
            return WD_ALIGN_PARAGRAPH.RIGHT
        return WD_ALIGN_PARAGRAPH.CENTER

    def _add_bingo_grid(self, doc: docx.Document, card_items: list[str]) -> None:
        """Add the centered Bingo matrix with fixed square cells."""
        rows = self.settings.rows
        cols = self.settings.cols
        table = doc.add_table(rows=rows, cols=cols)
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        table.allow_autofit = False
        set_table_layout_fixed(table)

        cell_size = self.metrics.grid_size
        cell_twips = int(cell_size * TWIPS_PER_INCH)

        for col_index in range(cols):
            table.columns[col_index].width = Inches(cell_size)

        for row_idx, row in enumerate(table.rows):
            set_row_height_exact(row, cell_twips)
            for col_idx, cell in enumerate(row.cells):
                cell.width = Inches(cell_size)
                set_column_width(cell, cell_twips)
                cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
                set_cell_margins(
                    cell,
                    top=self.settings.cell_padding_twips,
                    bottom=self.settings.cell_padding_twips,
                    left=self.settings.cell_padding_twips,
                    right=self.settings.cell_padding_twips,
                )

                fill = self.settings.cell_bg_dark if (row_idx + col_idx) % 2 else self.settings.cell_bg_light
                set_cell_appearance(
                    cell,
                    fill_hex=fill,
                    border_color_hex=self.settings.grid_border_color,
                    border_width_pt=self.settings.grid_border_width_pt,
                )

                index = row_idx * cols + col_idx
                item_text = card_items[index] if index < len(card_items) else ""
                paragraph = cell.paragraphs[0]
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                paragraph.paragraph_format.space_before = Pt(0)
                paragraph.paragraph_format.space_after = Pt(0)
                paragraph.paragraph_format.line_spacing = 0.95

                run = paragraph.add_run(item_text)
                run.font.name = self.settings.cell_font
                run.font.size = Pt(self.settings.cell_font_size_pt)
                run.font.color.rgb = rgb_from_hex(self.settings.cell_text_color)
